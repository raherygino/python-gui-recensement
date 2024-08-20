from PyQt5.QtWidgets import QTableWidgetItem, QFileDialog
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from openpyxl import Workbook
import os

from ...view import StudentDialog, AddStudentDialog
from ...models import StudentModel, Student, MarkModel, SubjectModel
from ...common import Utils, Function
from ...components import ConfirmDialog

class ActionPresenter:
    
    def __init__(self, parent):
        self.presenter = parent
        self.view = parent.view
        self.model:StudentModel = parent.model
        self.markModel:MarkModel = parent.modelMark
        self.subjectModel:SubjectModel = parent.modelSubject
        self.utils = Utils()
        self.func = Function()
        
    def studentByMatricule(self, matricule) -> Student:
        return self.model.fetch_item(matricule=matricule, promotion_id=self.presenter.promotionId)
    
    def addRow(self, table, nData:list):
        data = table.getData()
        table.setRowCount(len(data)+1)
        for col, value in enumerate(nData):
            item = QTableWidgetItem(value)
            if self.func.isFloat(value):
                item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            table.setItem(len(data), col, item)
        
    def showStudent(self, matricule):
        dialog = StudentDialog(self.view.parent.nParent)
        student = self.studentByMatricule(matricule)
        dialog.label.setText(f'{student.level} {student.matricule}\n{student.lastname} {student.firstname}')
        data = self.subjectModel.fetch_all(promotion_id=self.presenter.promotionId, level=student.level)
        dialog.table.setRowCount(len(data))
        sm = 0
        coefs = 0
        for i, sub in enumerate(data):
            mark = self.markModel.fetch_item(
                promotion_id=self.presenter.promotionId,
                student_id=student.id,
                subject_id=sub.id)
            mValue = self.func.strToFloat(str(mark.value)) if type(mark).__name__ == 'Marks' else '0'
            mValueCoef = self.func.strToFloat(str(mark.value*sub.coef)) if type(mark).__name__ == 'Marks' else '0'
            items = [sub.title, str(sub.coef), mValue ,mValueCoef]
            sm += self.func.toFloat(mValueCoef)
            coefs += sub.coef
            for col, itm in enumerate(items):
                item = QTableWidgetItem(itm)
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                if col != 0:
                    item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                dialog.table.setItem(i, col, item)
        self.addRow(dialog.table, ['TOTAL', '','',self.func.strToFloat(str(sm))])
        self.addRow(dialog.table, ['MOYENNE', '','',self.func.strToFloat(str(sm/coefs))])
        dialog.table.resizeColumnsToContents()
        dialog.exportWord.triggered.connect(lambda : self.exportStudent(student, dialog.table))
        dialog.exportExcel.triggered.connect(lambda : self.exportExcel(student, dialog.table))
        width =  50
        height = 500
        for i, item in enumerate(dialog.table.getHeaderLabels()):
            width += dialog.table.columnWidth(i)
        if width > 293:
            dialog.resize(width, height)
        dialog.exec()
        
    def exportStudent(self, student:Student, table):
        destination_path, _ = QFileDialog.getSaveFileName(self.view, "Exporter", "", "Text Files (*.docx);;All Files (*)")
        if destination_path:
            # Create a new Document
            doc = Document()
            doc.add_heading('Rélevé de Note', level=1)
            # Add some more paragraphs
            doc.add_paragraph(f'Nom et prénoms: {student.lastname} {student.firstname}')
            doc.add_paragraph(f'Matricule: {student.matricule}')
            doc.add_paragraph(f'Grade: {student.level}')
            data = [table.getHeaderLabels()]
            data.extend(table.getData())
            # Adding a table
            rows = len(data)
            cols = len(data[0])
            table = doc.add_table(rows=rows, cols=cols)
            # Set the style of the table (optional)
            table.style = 'Table Grid'
            # Set cell margins
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell.text = data[i][j]
                    for paragraph in cell.paragraphs:
                        if self.func.isFloat(data[i][j]):
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align text to the right
                        for run in paragraph.runs:
                            run.font.size = Pt(12)  # Set font size
            # Save the document
            doc.save(destination_path)
            os.startfile(destination_path)
            
            
    def exportExcel(self, student:Student, table):
        destination_path, _ = QFileDialog.getSaveFileName(self.view, "Exporter", "", "Text Files (*.xlsx);;All Files (*)")
        if destination_path:
            # Create a new Workbook and select the active worksheet
            wb = Workbook()
            ws = wb.active
            data = [
                ['Nom et prénoms', f'{student.lastname} {student.firstname}'],
                ['Matricule', student.matricule],
                ['Grade', student.level],
            ]
            data.extend([['']])
            data.extend([['']])
            data.extend([[sub for sub in table.getHeaderLabels()]])
            data.extend(table.getData())
            # Write data to the worksheet
            for row in data:
                ws.append(row)
                
            wb.save(destination_path)
            
    def deleteStudent(self, matricule):
        dialog = ConfirmDialog('Suppimer', 'Voulez-vous le suppimer vraiment?', self.view)
        dialog.setTitleBarVisible(False)
        if dialog.exec():
            self.model.delete_by(promotion_id=self.presenter.promotionId, matricule=matricule)
            self.presenter.setPromotionId(self.presenter.promotionId)
            self.utils.infoBarSuccess("Succès", "Suppression avec réussite", self.view)
            
    def deleteMultiple(self, matricules):
        dialog = ConfirmDialog('Suppimer', 'Voulez-vous le suppimer vraiment?', self.view)
        dialog.setTitleBarVisible(False)
        if dialog.exec():
            items = [{'promotion_id':self.presenter.promotionId, 'matricule':matricule} for matricule in matricules]
            self.model.delete_mutlitple(items)
            self.presenter.setPromotionId(self.presenter.promotionId)
            self.utils.infoBarSuccess("Succès", "Suppression avec réussite", self.view)
        
    def editStudent(self, matricule):
        dialog = AddStudentDialog(self.view.parent.nParent)
        student = self.studentByMatricule(matricule)  
        dialog.title.setText(f'Modifier {student.level} {student.lastname}')
        dialog.matriculeEdit.lineEdit.setText(str(student.matricule))
        dialog.matriculeEdit.lineEdit.setEnabled(False)
        dialog.lastnameEdit.lineEdit.setText(student.lastname)
        dialog.firstnameEdit.lineEdit.setText(student.firstname)
        dialog.genderEdit.combox.setCurrentIndex(1 if student.gender == 'F' else 0)
        dialog.gradeEdit.combox.setCurrentIndex(1 if student.level == 'EAP' else 0)
        dialog.yesBtn.clicked.connect(lambda: self.updateStudent(student, dialog))
        dialog.exec()
    
    def updateStudent(self,oldStudent:Student, dialog: AddStudentDialog):
        lastname  = dialog.lastnameEdit.lineEdit.text()
        firstname = dialog.firstnameEdit.lineEdit.text()
        grade     = dialog.gradeEdit.combox.currentText()
        gender    = dialog.genderEdit.combox.currentText()
        self.model.update_item(oldStudent.id, lastname=lastname, firstname=firstname, gender=gender, level=grade)
        self.presenter.setPromotionId(self.presenter.promotionId)
        self.utils.infoBarSuccess("Succès", "Mise à jour avec réussite", self.view)
        dialog.accept()