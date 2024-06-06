from PyQt5.QtCore import QPoint, QTimer
from PyQt5.QtGui import QCursor
from PyQt5.QtWidgets import QFileDialog
import os
from docx import Document
from qfluentwidgets import RoundMenu, Action, FluentIcon, MenuAnimationType, Dialog, MessageBox
from ..models import StudentModel, Mouvement, Student
from ..view.students.show_student_dialog import ShowStudentDialog
from ..view.students.new_movement_dialog import NewMouvementDialog
from ..view.students.new_student_dialog import NewStudentDialog
from ..common import Function

class MenuAction:
    
    def __init__(self, presenter) -> None:
        self.presenter = presenter
        self.view = presenter.view
        self.model:StudentModel = self.presenter.model
        #self.modelMove = self.presenter.modelMove
        self.ids = []
        self.func = Function()
        self.timer = QTimer()
        self.timer.setInterval(500)
        self.timer.timeout.connect(self.createMove)
        
    def createMove(self):
        self.modelMove.create(self.mouvementCreate)
        self.presenter.view.parent.nParent.refresh.emit(["mouvement"])
        self.timer.stop()
        
    def show(self, matricule):
        student = self.dataStudent(matricule)
        mouvements = self.modelMove.fetch_items_by_id(student.id)
        dataMouvements = self.refreshMove(mouvements)
        havingMove = len(dataMouvements) > 0
        day = self.modelMove.sum_by_with_id(student.id, "day")
        total = int(day) if day != "" else 0
        dataMouvements.append(["Total", "","","", total]) if total  > 0 else None
        
        dialog = ShowStudentDialog(self.presenter.mainView)
        dialog.label.setText(f'{student.level} {student.matricule}\n{student.lastname} {student.firstname}')
        dialog.ImageMessage.setVisible(not havingMove)
        dialog.message.setVisible(not havingMove)
        dialog.table.setVisible(havingMove)
        dialog.table.setData(dataMouvements)
        dialog.table.contextMenuEvent = lambda event, student=student, table=dialog.table: self.mouseRightClickTable(event, student, table)
        dialog.exportButton.clicked.connect(lambda: self.exportMouvement(student))  
        if student.matricule != 0:
            dialog.exec()
        else:
            self.func.errorSuccess("Elève introuvable", "Il n'existe plus dans la base de données", self.view)
    
    def exportMouvement(self, student):
        mouvements = self.modelMove.fetch_items_by_id(student.id)
        document = Document()
        document.add_heading(f'Informations', level=1)
        pData = f'Nom: {student.lastname}\n'
        pData += f'Prénoms: {student.firstname}\n'
        pData += f'Genre: {student.gender}\n'
        pData += f'Niveau: {student.level}\n'
        pData += f'Matricule: {student.matricule}\n'
        document.add_paragraph(pData)
        document.add_heading(f'Mouvements', level=1)
        
        # get table data -------------
        items = [[]]
        items.clear()
        day = "0"
        for mouv in mouvements:
            items.append([mouv.subType ,mouv.motif, mouv.date_start, mouv.date_end, mouv.day])
            if(len(mouv.day) != 0):
                day +=  "+"+mouv.day
        valDay = eval(day)
        if valDay > 0:
            items.append(["Total", "","","",str(valDay)])

        if len(items) == 0:
            document.add_paragraph("Aucun mouvement")
        else:
            # add table ------------------
            table = document.add_table(1, 5)
            table.style = "Table Grid"

            # populate header row --------
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Observations'
            heading_cells[1].text = 'Motifs'
            heading_cells[2].text = 'Date debut'
            heading_cells[3].text = 'Date fin'
            heading_cells[4].text = 'Nombre de jour'

            # add a data row for each item
            for item in items:
                cells = table.add_row().cells
                cells[0].text = str(item[0])
                cells[1].text = item[1]
                cells[2].text = item[2]
                cells[3].text = item[3]
                cells[4].text = item[4]
    
        # Save the document
        filename = f"{os.path.expanduser('~')}"
        fileName = self.dialogSaveFile("Exporter", filename, "Document Word (*.docx)")
        if fileName:
            document.save(fileName)
            os.startfile(fileName)  
        
    def dialogSaveFile(self, title:str, dir:str, typeFile:str):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getSaveFileName(self.view,title,dir,typeFile, options=options)
        return fileName
        
    def mouseRightClickTable(self, event, student, table):
        if len(table.selectedItems()) != 0 and table.currentRow() < len(self.ids):
            menu = RoundMenu(parent=self.view)
            menu.addAction(Action(FluentIcon.DELETE, 'Supprimer', triggered=lambda:self.delete_move(table.currentRow(), student, table)))
            self.posCur = QCursor().pos()
            cur_x = self.posCur.x()
            cur_y = self.posCur.y()

            menu.exec(QPoint(cur_x, cur_y), aniType=MenuAnimationType.FADE_IN_DROP_DOWN)
    
    def dialogConfirm(self):
        dialog = MessageBox(f"Supprimer", "Vous êtes sûr de vouloir supprimer?", self.presenter.mainView)
        dialog.yesButton.setText("Oui")
        dialog.cancelButton.setText("Non")
        return dialog
    
    def delete_move(self, index, student, table):
        w = Dialog("Supprimer", "Voulez-vous le supprimer vraiment", self.view.parent)
        w.setTitleBarVisible(False)
        if w.exec():
            self.modelMove.delete_by(idStudent=student.id, id=self.ids[index])
            dataMouvements = self.refreshMove(self.modelMove.fetch_items_by_id(student.id))
            if len(dataMouvements) > 0:
                total = int(self.modelMove.sum_by_with_id(student.id, "day"))
                if total > 0:
                    dataMouvements.append(["Total", "","","", total])
            
            table.setData(dataMouvements)
            self.presenter.view.parent.nParent.refresh.emit(["mouvement"])
            
    def update(self, matricule):
        oldStudent = self.model.fetch_item_by_cols(promotion_id=self.presenter.promotionId, matricule=matricule)
        dialog = NewStudentDialog(self.view)
        dialog.lastnameEdit.lineEdit(0).setText(oldStudent.lastname)
        dialog.firstnameEdit.lineEdit(0).setText(oldStudent.firstname)
        dialog.matriculeEdit.lineEdit(0).setText(str(oldStudent.matricule))
        if oldStudent.gender == "F":
            dialog.genderEdit.combox.setCurrentIndex(1)
        if oldStudent.level == "EAP":
            dialog.gradeEdit.combox.setCurrentIndex(1)
            
        dialog.yesButton.setEnabled(True)
        dialog.yesButton.setText("Mettre à jour")
        if dialog.exec():
            student = self.presenter.parent.dataStudentFromDialog(dialog)
            self.model.update_item(oldStudent.id, 
                              lastname=student.lastname, 
                              firstname=student.firstname,
                              gender=student.gender,
                              level=student.level,
                              matricule=student.matricule,
                              company=student.company,
                              section=student.section,
                              number=student.number)
            self.presenter.setPromotionId(self.presenter.promotionId)
            self.func.toastSuccess("Mise à jour", "Mise à jour d'élève avec réussite!", self.presenter.view)
        
    def mouvement(self, matricule):
        student = self.dataStudent(matricule)
        typeComp = self.presenter.typeCompModel.fetch_items_by_id(0)
        typeCompVal = [typeVal.name for typeVal in typeComp]
        dialog = NewMouvementDialog(self.view.parent)
        dialog.setMotifAutoComplete([item for item in self.presenter.fetchDataGroup(self.modelMove, key="motif")])
        dialog.typeEdit.combox.addItems(typeCompVal)
        self.typeMoveChanged(0, typeCompVal, dialog)
        dialog.typeEdit.combox.currentIndexChanged.connect(lambda index, data=typeCompVal: self.typeMoveChanged(index, data, dialog))
        dialog.subTitle.setText(f'{student.level} {student.lastname} {student.firstname}')
        if dialog.exec():
            matricule = str(student.matricule)
            day = dialog.dayEdit.text(0)
            self.mouvementCreate = Mouvement(
                idStudent=student.id,
                promotion_id=self.presenter.promotionId,
                level=student.level,
                matricule=student.matricule,
                company=matricule[0],
                section=matricule[1],
                gender=student.gender,
                student=str(student),
                type=dialog.typeEdit.combox.currentText(),
                subType=dialog.subTypeEdit.combox.currentText(),
                motif=dialog.motifEdit.lineEdit(0).text(),
                date_start=dialog.dateEdit.date.text(),
                date_end=dialog.dateEdit2.date.text(),
                day= day if day != "" else 0
            )
            self.timer.start()
            
    def typeMoveChanged(self, index, data, dialog):
        Comp = self.presenter.compModel.fetch_items_by_col(self.presenter.promotionId, comp_type=data[index])
        compVal = [cVal.name for cVal in Comp]
        dialog.subTypeEdit.combox.clear()
        dialog.subTypeEdit.combox.addItems(compVal)
        
    def delete(self, matricule):
        if self.dialogConfirm().exec():
            student = self.model.fetch_all(promotion_id=self.presenter.promotionId, matricule=matricule)[0]
            self.model.delete_item(student.id)
            self.presenter.setPromotionId(self.presenter.promotionId)
            self.presenter.view.parent.nParent.refresh.emit(["mouvement"])
            
    def deleteMove(self, **kwargs):
        if self.dialogConfirm().exec():
            kwargs["promotion_id"] = self.presenter.promotionId
            self.modelMove.delete_by(**kwargs)
            self.presenter.view.parent.nParent.refresh.emit(["mouvement"])
            
    def deleteMultipleMove(self, items:list):
        if self.dialogConfirm().exec():
            self.modelMove.delete_mutlitple(items)
            self.presenter.view.parent.nParent.refresh.emit(["mouvement"])
    
    def dataStudent(self, matricule):
        data = self.model.fetch_all(promotion_id=self.presenter.promotionId, matricule=matricule)
        return data[0] if len(data) != 0 else Student()
    
    def refreshMove(self, mouvements):
        dataMouvements = []
        self.ids.clear()
        for mouvement in mouvements:
            self.ids.append(mouvement.id)
            dataMouvements.append([
                mouvement.subType,
                mouvement.motif,
                mouvement.date_start,
                mouvement.date_end,
                mouvement.day
            ])
        return dataMouvements