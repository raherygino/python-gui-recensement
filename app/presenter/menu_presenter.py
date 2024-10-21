from PyQt5.QtCore import QDate, QPoint
from PyQt5.QtGui import QCursor
from PyQt5.QtWidgets import QFileDialog
from ..components import ConfirmDialog
from ..models import Believer, BelieverModel
from ..view import ViewBelieverDialog, AddBelieverInterface, WithFamilyDialog
from qfluentwidgets import MessageDialog, RoundMenu, Action, MenuAnimationType, FluentIcon
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
import os

class MenuAction:
    
    def __init__(self, presenter) -> None:
        self.presenter = presenter
        self.view = presenter.view
        self.model: BelieverModel= self.presenter.model
        
    def show(self, itemId):
        believer: Believer = self.model.fetch_item(id=itemId)
        row1 = [
            ["Anarana", believer.lastname],
            ["Fanampiny", believer.firstname],
            ["Adiresy", believer.address],
            ["Faritra", believer.region],
            ["Diakonina miandraikitra", believer.diacon]
            ]
        
        row2 = [
            ["Daty sy toerana nahaterahana", f'{believer.birthday} {believer.birthplace}'],
            ["Anaran'i Ray", believer.name_father],
            ["Anaran'i Reny", believer.name_mother],
            ["Daty ny batisa", believer.date_of_baptism],
            ]
        
        row3 = [
            ["Toerana ny batisa", believer.place_of_baptism],
            ["Daty nahampandray", believer.date_of_recipient],
            ["Toerana nahampandray", believer.place_of_recipient],
            ["Laharana ny mpandray", believer.number_recipient],
            ["Laharan'ny finday", believer.phone],
            ]
        
        row4 = [
            ["Andraikitra", believer.responsibility],
                ["Asa", believer.work],
                ["Fanamarihana", believer.obs],
            ]
            
        allRows =  [row1, row2, row3, row4]
        
        dialog = ViewBelieverDialog(self.view)
        dialog.table.setHorizontalHeaderLabels(self.presenter.labelsFamily)
            
        for i, row in enumerate(allRows):
            for j, column in enumerate(row):
                dialog.addLabelValue(column[0], column[1], i,j)
                
        dialog.addNewLabelValue("Sampana", believer.dept_work)
                    
        data = self.model.fetch_all(id_conjoint=itemId)
        for value in self.model.fetch_all(id_father=itemId):
            data.append(value)
                
        self.presenter.setData(dialog.table, data)
        dialog.table.contextMenuEvent = lambda e, dialog = dialog, data=data: self.rightClickTable(e, data, dialog)
        dialog.yesBtn.setText("Avoaka WORD")
        dialog.cancelBtn.setText("Miverina")
        if dialog.exec():
            wDialog = WithFamilyDialog(self.view)
            if wDialog.exec():
                options = QFileDialog.Options()
                fileName, _ = QFileDialog.getSaveFileName(self.presenter.addView.nParent,"Export File", "","Word Files (*.docx);;All Files (*)", options=options)
                if fileName:
                    # Create a new Document
                    doc = Document()
                    doc.styles['Normal'].paragraph_format.line_spacing = Pt(12)  # Set line spacing to 20 points (adjust as needed)

                    # Set page orientation to landscape
                    section = doc.sections[0]
                    
                    # Set margins
                    section = doc.sections[0]
                    section.left_margin = Inches(0.2)   # 1 inch margin on the left
                    section.right_margin = Inches(0.2)  # 1 inch margin on the right
                    section.top_margin = Inches(0.2)     # 1 inch margin at the top
                    section.bottom_margin = Inches(0.2)  # 1 inch margin at the bottom
                    
                    new_width, new_height = section.page_height, section.page_width
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = new_width
                    section.page_height = new_height
                    # Add a heading
                    doc.add_heading('Loham-pianakaviana', level=1)
                    # Add some paragraphs
                    '''doc.add_paragraph('This is a sample paragraph.')
                    doc.add_paragraph('This is another paragraph.')

                    # Add a numbered list
                    doc.add_paragraph('Numbered List:', style='List Number')
                    for i in range(1, 4):
                        doc.add_paragraph(f'Item {i}', style='List Number')'''
                    # Data for the table
                    dataBeliever = [
                        [f'Anarana: \n{believer.lastname}', 
                        f'Fanampiny: \n{believer.firstname}', 
                        f'Adiresy: \n{believer.address}',
                        f'Faritra: \n{believer.region}'],
                        
                        [f'Diakonina miandraikitra: \n{believer.diacon}', 
                        f'Daty sy toerana nahaterahana: \n{believer.birthday} {believer.birthplace}', 
                        f'Anaran\'i Ray: \n{believer.name_father}',
                        f'Anaran\'i Reny: \n{believer.name_mother}'],
                        
                        [f'Daty ny batisa: \n{believer.date_of_baptism}', 
                        f'Toerana ny batisa: \n{believer.birthday}', 
                        f'Daty nahampandray: \n{believer.date_of_recipient}',
                        f'Toerana nahampandray: \n{believer.place_of_recipient}'],
                        
                        [f'Larahana nahampandray: \n{believer.number_recipient}', 
                        f'Laharan\'ny finday: \n{believer.phone}', 
                        f'Sampana na/sy sampan\'asa: \n{believer.dept_work}',
                        f'Andraikitra: \n{believer.responsibility}'],
                        
                        [f'Fanamarihana: \n{believer.obs}', 
                        '', 
                        '',
                        ''],
                    ]
                    # Add a table
                    table = doc.add_table(rows=len(dataBeliever), cols=len(dataBeliever[0]))
                    for i, row in enumerate(dataBeliever):
                        for j, value in enumerate(row):
                            cell = table.cell(i, j)
                            for k, headline in enumerate(value.split('\n')):
                                p = cell.add_paragraph()
                                run = p.add_run(headline)
                                if k  == 0:
                                    run.bold = True
                                    run.underline = True
                                
                    # Add a heading 
                    if wDialog.combox.currentText() == "Eny":
                        doc.add_heading('Fianakaviana', level=1)
                        
                        dataFamily = [
                            ["Anarana","Fanampiny","Lahy sa vavy","Amin'ny fianakaviana",
                            "Daty sy toerana nahaterahana","Batisa","Daty sy toerana maha mpandray",
                            "Laharana karatra mpandray","Sampana sy/na Sampan'asa", "Asa", "Fanamarihana"]
                        ]
                        for blv in data:
                            blv: Believer = blv
                            dataFamily.append([
                                blv.lastname,
                                blv.firstname,
                                blv.gender,
                                blv.pos_family,
                                f'{blv.birthday} {blv.birthplace}',
                                f'{blv.date_of_baptism} {blv.place_of_baptism}',
                                f'{blv.date_of_recipient} {blv.place_of_recipient}',
                                blv.number_recipient,
                                blv.dept_work,
                                blv.work,
                                blv.obs
                            ])
                            
                        tablefamily = doc.add_table(rows=len(dataFamily), cols=len(dataFamily[0]))
                        tablefamily.style = 'Table Grid'
                        
                        for i, row in enumerate(dataFamily):
                            for j, value in enumerate(row):
                                tablefamily.cell(i, j).text = value

                    # Save the document
                    doc.save(fileName)
                    os.startfile(fileName)

        
    def rightClickTable(self, event, data, dialog):
        selectedItems = dialog.table.selectedItems()
        selectedIndex = dialog.table.selectedIndexes()
        if len(selectedItems) != 0:
            if selectedItems[3].text() == "Lahy":
                for i, index in enumerate(selectedIndex):
                    if i == 0:
                        row = index.row()
                        blv = data[row]
                        if blv.is_leader == 0:
                            menu = RoundMenu(parent=self.view)
                            menu.addAction(
                            Action(
                                FluentIcon.PEOPLE, 
                                'Loham-pianakaviana hafa', 
                                triggered= lambda: self.newLeaderFamily(blv, dialog)))
            
                            self.posCur = QCursor().pos()
                            cur_x = self.posCur.x()
                            cur_y = self.posCur.y()
                            menu.exec(QPoint(cur_x, cur_y), aniType=MenuAnimationType.FADE_IN_DROP_DOWN)
                
    def newLeaderFamily(self, item, dialog):
        self.update(item.id)
        self.presenter.isNewLead = True
        dialog.hide()
        
    def strToQDate(self, strDate: str):
        date = strDate.split("/")
        dateOut = QDate(2000,1,1)
        if len(date) == 3:
            dateOut = QDate(int(date[2]), int(date[1]), int(date[0]))
        return dateOut
        
    def update(self, item):
        self.presenter.addView.nParent.stackedWidget.setCurrentWidget(self.presenter.addView)
        self.presenter.idEdit = int(item)
        self.presenter.isNewLead = False
        view : AddBelieverInterface = self.presenter.addView
        believer : Believer = self.model.fetch_item(id=item)
        idFather = believer.id_father
        idMother = believer.id_mother
        if idFather != 0:
            father : Believer = self.model.fetch_item(id=idFather)
            view.nameFatherEdit.lineEdit.setText(f'{father.lastname} {father.firstname}')
        else:
            view.nameFatherEdit.lineEdit.setText(believer.name_father)
        if idMother != 0:
            mother : Believer = self.model.fetch_item(id=idMother)
            view.nameMotherEdit.lineEdit.setText(f'{mother.lastname} {mother.firstname}')
        else:
            view.nameMotherEdit.lineEdit.setText(believer.name_mother)
        view.lastnameEdit.lineEdit.setText(believer.lastname)
        view.firstnameEdit.lineEdit.setText(believer.firstname)
        view.addressEdit.lineEdit.setText(believer.address)
        view.regionEdit.combox.setCurrentText(believer.region)
        view.diaconEdit.combox.setCurrentText(believer.diacon)
        view.birthdayEdit.lineEdit.setDate(self.strToQDate(believer.birthday))
        view.birthplaceEdit.lineEdit.setText(believer.birthplace)
        view.baptismDateEdit.lineEdit.setDate(self.strToQDate(believer.date_of_baptism))
        view.baptismPlaceEdit.lineEdit.setText(believer.place_of_baptism)
        view.recipientDateEdit.lineEdit.setDate(self.strToQDate(believer.date_of_recipient))
        view.recipientPlaceEdit.lineEdit.setText(believer.place_of_recipient)
        view.recipientNumberEdit.lineEdit.setText(believer.number_recipient)
        view.phoneEdit.lineEdit.setText(believer.phone)
        view.deptWorkCheck.check(believer.dept_work)
        view.responsibilityEdit.lineEdit.setText(believer.responsibility)
        view.workEdit.lineEdit.setText(believer.work)
        view.obsEdit.lineEdit.setText(believer.obs)
        
        data = self.model.fetch_all(id_conjoint=item)
        for value in self.model.fetch_all(id_father=item):
            data.append(value)
        self.presenter.family = data
        self.presenter.setData(view.familyTableView, self.presenter.family)
        self.presenter.addView.btnAdd.setText("Ovaina")
        
    def confirmDelete(self, item):
        #dialog = MessageDialog('Fafana', 'Fafana marina ve?', self.view.nParent)
        dialog = ConfirmDialog('Fafana', 'Fafana marina ve?',  self.view.nParent)
        dialog.yesBtn.setText("Eny")
        dialog.cancelBtn.setText("Tsia")
        #dialog.yesButton.clicked.connect(lambda: self.delete(item))
        if dialog.exec_():
            nItem = self.model.fetch_item(id=item)
            if nItem.id_father == 0:
                self.model.delete_item(item)
            else:
                self.model.update_item(item, is_leader="0", gender="Lahy")
                self.model.delete_by(id_conjoint=nItem.id)
                self.model.delete_by(id_father=nItem.id)
            self.presenter.fetchData(self.model.fetch_all(**self.presenter.query))
            
            
        
    def delete(self, item):
        self.model.delete_item(item)
        self.presenter.fetchData(self.model.fetch_all(**self.presenter.query))
        #self.presenter.refresh.emit()